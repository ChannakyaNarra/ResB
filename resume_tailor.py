import os
import shutil
import uuid
import asyncio
import re
import logging
from openai import OpenAI, OpenAIError
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pdf2docx import Converter
from docx import Document
from docx.shared import Pt
from docx2pdf import convert
from dotenv import load_dotenv

# --- INITIALIZATION AND CONFIGURATION ---

# Load environment variables from .env file
load_dotenv()
GITHUB_TOKEN = os.getenv("GITHUB_TOKEN")
if not GITHUB_TOKEN:
    raise RuntimeError("FATAL: Missing GITHUB_TOKEN in environment variables.")

# Initialize the AI client for GitHub Models
try:
    client = OpenAI(
        base_url="https://models.github.ai/inference",
        api_key=GITHUB_TOKEN,
    )
except Exception as e:
    raise RuntimeError(f"Failed to initialize OpenAI client: {e}")

# Set up logging to monitor application activity and errors
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger("uvicorn.error")

# Initialize FastAPI app
app = FastAPI(title="Resume Tailoring API v3")

# Add CORS middleware for frontend integration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Directory for storing uploaded and generated files
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# --- CORE AI AND DOCUMENT PROCESSING LOGIC ---

# Use a semaphore to limit concurrent API requests to prevent rate-limiting.
# A value of 3-5 is a safe starting point.
API_CONCURRENCY_LIMIT = 4
semaphore = asyncio.Semaphore(API_CONCURRENCY_LIMIT)


async def tailor_bullet(section_name: str, job_description: str, original_text: str) -> str:
    """
    Asynchronously calls the AI model, respecting the semaphore to limit concurrency.
    """
    async with semaphore:
        logger.info(f"Semaphore acquired. Tailoring bullet for section: {section_name}")
        prompt = (
            f"Rewrite this single resume bullet point from the '{section_name}' section to be more impactful and align "
            f"closely with the following job description. **Crucially, you must keep the word count of your rewritten "
            f"bullet point as close as possible to the original to preserve the document's layout.** "
            f"Return only the single, rewritten bullet point text, without any prefixes like '•' or '-'.\n\n"
            f"Job Description: \"{job_description}\"\n\n"
            f"Original Bullet Point: \"{original_text}\""
        )
        messages = [
            {"role": "system",
             "content": "You are an expert resume writer. Your task is to refine resume content to be concise and perfectly tailored to a job description, while strictly maintaining the original's approximate word count to preserve document formatting."},
            {"role": "user", "content": prompt}
        ]
        try:
            loop = asyncio.get_event_loop()
            response = await loop.run_in_executor(
                None,
                lambda: client.chat.completions.create(
                    model="openai/gpt-4o-mini",
                    messages=messages,
                    temperature=0.6,
                    max_tokens=200,
                    top_p=1.0
                )
            )
            rewritten_text = response.choices[0].message.content.strip()
            return re.sub(r'^[•*-]\s*', '', rewritten_text)

        except OpenAIError as e:
            if isinstance(e, OpenAIError) and "Too Many Requests" in str(e):
                logger.warning(
                    "Rate limit hit, will be retried by the client library. Consider lowering API_CONCURRENCY_LIMIT if this persists.")
            logger.error(f"GitHub AI API error: {e}", exc_info=True)
            raise HTTPException(status_code=502, detail=f"An error occurred with the AI service: {e}")
        except Exception as e:
            logger.error(f"Unexpected error in tailor_bullet: {e}", exc_info=True)
            return original_text  # Fallback to original text


def get_paragraph_formatting(para):
    """Captures all formatting from a paragraph's first run."""
    if not para.runs:
        return {}

    first_run = para.runs[0]
    font = first_run.font
    return {
        'style': para.style.name,
        'font_name': font.name,
        'font_size': font.size,
        'bold': font.bold,
        'italic': font.italic,
        'underline': font.underline,
        'color': font.color.rgb if font.color and font.color.rgb else None,
        'paragraph_format': {
            'alignment': para.alignment,
            'left_indent': para.paragraph_format.left_indent,
            'right_indent': para.paragraph_format.right_indent,
            'first_line_indent': para.paragraph_format.first_line_indent,
            'space_before': para.paragraph_format.space_before,
            'space_after': para.paragraph_format.space_after,
            'line_spacing': para.paragraph_format.line_spacing,
        }
    }


def apply_paragraph_formatting(para, new_text, formatting):
    """Applies captured formatting to a new run in a cleared paragraph."""
    # Clear existing content by removing all runs
    for run in list(para.runs):
        p = para._p
        p.remove(run._r)

    # Add the new text in a single run
    run = para.add_run(new_text)

    # Apply font formatting
    font = run.font
    font.name = formatting.get('font_name')
    font.size = formatting.get('font_size')
    font.bold = formatting.get('bold')
    font.italic = formatting.get('italic')
    font.underline = formatting.get('underline')
    if formatting.get('color'):
        font.color.rgb = formatting['color']

    # Apply paragraph formatting
    p_fmt = para.paragraph_format
    fmt_p_info = formatting.get('paragraph_format', {})
    if fmt_p_info.get('alignment') is not None:
        p_fmt.alignment = fmt_p_info['alignment']
    p_fmt.left_indent = fmt_p_info.get('left_indent')
    p_fmt.right_indent = fmt_p_info.get('right_indent')
    p_fmt.first_line_indent = fmt_p_info.get('first_line_indent')
    p_fmt.space_before = fmt_p_info.get('space_before')
    p_fmt.space_after = fmt_p_info.get('space_after')
    p_fmt.line_spacing = fmt_p_info.get('line_spacing')


def is_section_header(paragraph) -> bool:
    """Flexibly checks if a paragraph is a section header."""
    text = paragraph.text.strip()
    if not text:
        return False
    header_pattern = re.compile(
        r'^(PROFESSIONAL EXPERIENCE|WORK HISTORY|EXPERIENCE|EDUCATION|SKILLS|OBJECTIVE|SUMMARY|TECHNICAL SKILLS):?$',
        re.IGNORECASE)
    is_bold = any(run.bold for run in paragraph.runs)
    is_all_caps = text.isupper() and len(text.split()) < 5
    return bool(is_bold or is_all_caps or header_pattern.match(text))


def is_bullet_point(paragraph) -> bool:
    """Checks if a paragraph is a bullet point."""
    style_is_list = paragraph.style.name.lower().startswith('list')
    text_is_bullet = paragraph.text.strip().startswith(("-", "•", "*"))
    return style_is_list or text_is_bullet


async def process_docx(docx_path: str, job_description: str) -> str:
    """
    Processes the DOCX file, preserving formatting and concurrently tailoring bullets.
    """
    doc = Document(docx_path)
    current_section = "Unknown"
    tasks = []
    paragraphs_to_process = []

    for para in doc.paragraphs:
        if is_section_header(para):
            current_section = para.text.strip().capitalize()
            continue

        text = para.text.strip()
        if text and (is_bullet_point(para) or "Skills" in current_section):
            # Strip the bullet character itself for a cleaner prompt
            original_text_for_ai = re.sub(r'^[•*-]\s*', '', text)
            formatting = get_paragraph_formatting(para)
            paragraphs_to_process.append({'para': para, 'text': original_text_for_ai, 'format': formatting,
                                          'prefix': text.split(' ')[0] + ' ' if text.startswith(
                                              ("-", "•", "*")) else ''})

    for item in paragraphs_to_process:
        # Create the task, which will wait for the semaphore before running
        task = tailor_bullet(current_section, job_description, item['text'])
        tasks.append(task)

    if tasks:
        logger.info(f"Processing {len(tasks)} bullets with a concurrency limit of {API_CONCURRENCY_LIMIT}...")
        tailored_results = await asyncio.gather(*tasks, return_exceptions=True)
        logger.info("AI tailoring complete. Applying changes to document.")

        # Check for exceptions returned by asyncio.gather
        for i, result in enumerate(tailored_results):
            if isinstance(result, Exception):
                logger.error(
                    f"Task for bullet {i + 1} failed with exception: {result}. The original text will be kept.")
                # We can get the original text from our list to ensure the document is still generated
                tailored_results[i] = paragraphs_to_process[i]['text']

        for i, item in enumerate(paragraphs_to_process):
            full_new_text = f"{item['prefix']}{tailored_results[i]}"
            apply_paragraph_formatting(item['para'], full_new_text, item['format'])

    tailored_docx_path = docx_path.replace(".docx", "_tailored.docx")
    doc.save(tailored_docx_path)
    return tailored_docx_path


async def process_resume(resume_path: str, job_description: str) -> tuple[str, str]:
    """Main processing pipeline: PDF -> DOCX -> Tailor -> PDF."""
    try:
        temp_docx_path = resume_path.replace(".pdf", ".docx")
        logger.info(f"Converting {resume_path} to DOCX...")
        cv = Converter(resume_path)
        cv.convert(temp_docx_path, start=0, end=None)
        cv.close()
        logger.info("Conversion to DOCX successful.")

        logger.info("Processing DOCX with AI (formatting-aware, rate-limited)...")
        tailored_docx_path = await process_docx(temp_docx_path, job_description)
        logger.info("DOCX processing complete.")

        tailored_pdf_path = tailored_docx_path.replace(".docx", ".pdf")
        logger.info(f"Converting {tailored_docx_path} back to PDF...")
        convert(tailored_docx_path, tailored_pdf_path)
        logger.info("Conversion to PDF successful.")

        os.remove(resume_path)
        os.remove(temp_docx_path)

        return tailored_pdf_path, tailored_docx_path
    except Exception as e:
        logger.error(f"Error in process_resume pipeline: {e}", exc_info=True)
        # Clean up failed intermediate files
        if 'temp_docx_path' in locals() and os.path.exists(temp_docx_path):
            os.remove(temp_docx_path)
        raise HTTPException(status_code=500, detail=f"An internal error occurred during resume processing: {e}")


# --- API ENDPOINTS ---

@app.post("/tailor-resume")
async def tailor_resume_endpoint(
        job_description: str = Form(...),
        resume: UploadFile = File(...)
):
    """Endpoint to upload a resume PDF and job description."""
    secure_suffix = str(uuid.uuid4())
    original_filename = os.path.splitext(resume.filename)[0].replace(" ", "_")
    pdf_path = os.path.join(UPLOAD_DIR, f"{original_filename}_{secure_suffix}.pdf")

    logger.info(f"Receiving file '{resume.filename}'. Saving as '{pdf_path}'.")
    try:
        with open(pdf_path, "wb") as f:
            shutil.copyfileobj(resume.file, f)

        tailored_pdf_path, tailored_docx_path = await process_resume(pdf_path, job_description)

        return {
            "message": "Resume tailored successfully while preserving format!",
            "pdf_download_url": f"/download/{os.path.basename(tailored_pdf_path)}",
            "docx_download_url": f"/download/{os.path.basename(tailored_docx_path)}",
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in tailor_resume_endpoint: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to process the resume.")


@app.get("/download/{filename}")
async def download_file(filename: str, background_tasks: BackgroundTasks):
    """Serves a generated file for download and cleans it up afterward."""
    file_path = os.path.join(UPLOAD_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found.")

    background_tasks.add_task(os.remove, file_path)

    return FileResponse(path=file_path, filename=filename, media_type='application/octet-stream')


@app.get("/", summary="Health Check")
async def health_check():
    """Health check endpoint."""
    return {"status": "OK", "message": "Resume Tailoring API v3 is running."}

